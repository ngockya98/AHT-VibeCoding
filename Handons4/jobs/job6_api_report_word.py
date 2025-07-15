"""
Job 6: Lấy yêu cầu từ API khách hàng → Sinh báo cáo Word
Tools: requests, python-docx
"""
import requests
from docx import Document

url = 'https://api.example.com/report'
data = requests.get(url).json()
doc = Document()
doc.add_heading('Báo cáo', 0)
for k, v in data.items():
    doc.add_paragraph(f'{k}: {v}')
doc.save('report.docx')
